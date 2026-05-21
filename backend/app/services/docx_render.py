"""Render markdown content (CVs, cover letters) to a Word document.

Uses python-docx. Covers the constructs the LLM emits per the prompts in
`app/services/generator.py`:

- ATX headers `#`, `##`, `###` → Heading 1/2/3
- Paragraphs with inline **bold** and *italic*
- Bulleted lists (`- item`)
- Numbered lists (`1. item`)
- Empty lines as paragraph breaks
- Horizontal rules (`---`) ignored

Anything fancier (tables, images, links) is rendered as plain text — fine for
CVs and cover letters that are mostly running prose.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
# Match **bold** and *italic* in order (greedy bold first, then italic).
_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _add_inline(paragraph, text: str) -> None:
    """Append text to `paragraph`, parsing **bold** / *italic*."""
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text: str) -> bytes:
    """Render markdown to a `.docx` file and return its bytes."""
    doc = Document()

    # Base style — Calibri 11 is Word default; we tighten line spacing.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line):
            i += 1
            continue

        # Heading
        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Bulleted list (consume contiguous bullets)
        if _BULLET_RE.match(line):
            while i < len(lines) and _BULLET_RE.match(lines[i].rstrip()):
                m = _BULLET_RE.match(lines[i].rstrip())
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, m.group(1) if m else "")
                i += 1
            continue

        # Numbered list
        if _NUMBERED_RE.match(line):
            while i < len(lines) and _NUMBERED_RE.match(lines[i].rstrip()):
                m = _NUMBERED_RE.match(lines[i].rstrip())
                p = doc.add_paragraph(style="List Number")
                _add_inline(p, m.group(1) if m else "")
                i += 1
            continue

        # Regular paragraph (may span multiple consecutive non-empty lines)
        para_lines: list[str] = []
        while i < len(lines):
            ln = lines[i].rstrip()
            if not ln.strip():
                break
            if (
                _HEADING_RE.match(ln)
                or _BULLET_RE.match(ln)
                or _NUMBERED_RE.match(ln)
            ):
                break
            para_lines.append(ln)
            i += 1
        if para_lines:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            _add_inline(p, " ".join(para_lines))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
